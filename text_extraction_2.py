import os
import re
import pdfplumber
import fitz  # PyMuPDF
import PyPDF2
from docx import Document

# Directory paths
data_dir = 'Dataset'
output_dir = 'Extracted_Data'

# Ensure output directory exists
os.makedirs(output_dir, exist_ok=True)

# Function to clean text
def clean_text(text):
    text = re.sub(r'\n+', '\n', text)  # Remove multiple newlines
    text = re.sub(r'\s+', ' ', text)   # Remove extra spaces
    return text.strip()

# Function to extract text and tables from PDF
def process_pdf(file_path):
    text = ''
    tables = []

    # Try using pdfplumber first
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            text += page.extract_text() or ''
            page_tables = page.extract_tables()
            if page_tables:
                tables.extend(page_tables)

    # If pdfplumber fails to extract meaningful text, fallback to PyMuPDF
    if not text.strip():
        with fitz.open(file_path) as pdf:
            for page in pdf:
                text += page.get_text()

    return clean_text(text), tables

# Function to extract text and tables from Word documents
def process_word(file_path):
    text = ''
    tables = []

    doc = Document(file_path)
    for para in doc.paragraphs:
        text += para.text + '\n'

    for table in doc.tables:
        tables.append([[cell.text for cell in row.cells] for row in table.rows])

    return clean_text(text), tables

# Process files in the Dataset directory
def process_files(data_dir, output_dir):
    for file in os.listdir(data_dir):
        file_path = os.path.join(data_dir, file)

        if file.endswith('.pdf'):
            output_path = os.path.join(output_dir, file.replace('.pdf', '.txt'))
            text, tables = process_pdf(file_path)

            with open(output_path, 'w', encoding='utf-8') as text_file:
                text_file.write(text if text else 'No text extracted.')

            if tables:
                table_output_path = os.path.join(output_dir, file.replace('.pdf', '_tables.txt'))
                with open(table_output_path, 'w', encoding='utf-8') as table_file:
                    for i, table in enumerate(tables):
                        table_file.write(f"\nTable {i + 1}:\n")
                        for row in table:
                            row = [str(item) if item is not None else '' for item in row]
                            table_file.write("\t".join(row) + "\n")

        elif file.endswith('.docx'):
            output_path = os.path.join(output_dir, file.replace('.docx', '.txt'))
            try:
                text, tables = process_word(file_path)

                with open(output_path, 'w', encoding='utf-8') as text_file:
                    text_file.write(text if text else 'No text extracted.')

                if tables:
                    table_output_path = os.path.join(output_dir, file.replace('.docx', '_tables.txt'))
                    with open(table_output_path, 'w', encoding='utf-8') as table_file:
                        for i, table in enumerate(tables):
                            table_file.write(f"\nTable {i + 1}:\n")
                            for row in table:
                                row = [str(item) if item is not None else '' for item in row]
                                table_file.write("\t".join(row) + "\n")

            except Exception as e:
                print(f"Error processing file {file_path}: {e}")

# Run the process
process_files(data_dir, output_dir)

print("Data Extraction Done!")
