import os
import re
import pdfplumber
from docx import Document
import fitz  # PyMuPDF

# Directory path
data_dir = 'Dataset'
output_dir = 'Extracted_Data_3'
unable_to_process_log = 'UnableToProcess.txt'

# Ensure output directory exists
os.makedirs(output_dir, exist_ok=True)

# Function to clean text
def clean_text(text):
    text = re.sub(r'\n+', '\n', text)  # Remove multiple newlines
    text = re.sub(r'\s+', ' ', text)   # Remove extra spaces
    return text.strip()

# Log for files that cannot be processed
unable_to_process = []

# Process PDF files
def process_pdf(file_path):
    output_path = os.path.join(output_dir, os.path.basename(file_path).replace('.pdf', '.txt'))
    text = ''
    tables = []

    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                # Extract text
                page_text = page.extract_text() or ''
                text += page_text

                # Extract tables
                page_tables = page.extract_tables()
                if page_tables:
                    tables.extend(page_tables)

        cleaned_text = clean_text(text)

        # Save cleaned text only if not empty
        if cleaned_text:
            with open(output_path, 'w', encoding='utf-8') as text_file:
                text_file.write(cleaned_text)

            # Save tables if any
            if tables:
                table_output_path = os.path.join(output_dir, os.path.basename(file_path).replace('.pdf', '_tables.txt'))
                with open(table_output_path, 'w', encoding='utf-8') as table_file:
                    for i, table in enumerate(tables):
                        table_file.write(f"\nTable {i + 1}:\n")
                        for row in table:
                            row = [str(item) if item is not None else '' for item in row]
                            table_file.write("\t".join(row) + "\n")
    except Exception as e:
        unable_to_process.append(f"{file_path} - PDF processing error: {e}")

# Process Word documents
def process_word(file_path):
    output_path = os.path.join(output_dir, os.path.basename(file_path).replace('.docx', '.txt'))
    text = ''
    tables = []

    try:
        doc = Document(file_path)
        for para in doc.paragraphs:
            text += para.text + '\n'

        for table in doc.tables:
            tables.append([[cell.text for cell in row.cells] for row in table.rows])

        cleaned_text = clean_text(text)

        # Save cleaned text only if not empty
        if cleaned_text:
            with open(output_path, 'w', encoding='utf-8') as text_file:
                text_file.write(cleaned_text)

            # Save tables if any
            if tables:
                table_output_path = os.path.join(output_dir, os.path.basename(file_path).replace('.docx', '_tables.txt'))
                with open(table_output_path, 'w', encoding='utf-8') as table_file:
                    for i, table in enumerate(tables):
                        table_file.write(f"\nTable {i + 1}:\n")
                        for row in table:
                            row = [str(item) if item is not None else '' for item in row]
                            table_file.write("\t".join(row) + "\n")
    except Exception as e:
        unable_to_process.append(f"{file_path} - Word processing error: {e}")

# Process files in the Dataset directory
def process_files(data_dir):
    for file in os.listdir(data_dir):
        file_path = os.path.join(data_dir, file)

        if file.endswith('.pdf'):
            process_pdf(file_path)
        elif file.endswith('.docx'):
            process_word(file_path)
        else:
            unable_to_process.append(f"{file_path} - Unsupported file extension")

# Run the process
process_files(data_dir)

# Log unable to process files
if unable_to_process:
    with open(unable_to_process_log, 'w', encoding='utf-8') as log_file:
        for error in unable_to_process:
            log_file.write(error + '\n')

print("Data Extraction Done!")
